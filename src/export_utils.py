"""Export utilities for Markdown, Word, and PDF formats."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone


def _sanitize_filename(name: str) -> str:
    """Build a safe base filename without extension."""
    cleaned = re.sub(r"[^\w\-]+", "_", name.strip(), flags=re.UNICODE)
    return cleaned[:80] or "documind_export"


def _pdf_safe_text(text: str) -> str:
    """Encode text for Helvetica core fonts (Latin-1)."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _break_long_words(text: str, max_word_len: int = 50) -> str:
    """Insert breaks in very long tokens so fpdf2 can wrap lines."""
    parts: list[str] = []
    for word in text.split():
        while len(word) > max_word_len:
            parts.append(word[:max_word_len])
            word = word[max_word_len:]
        parts.append(word)
    return " ".join(parts)


def _pdf_write_line(pdf, text: str, line_height: float, font_size: int, bold: bool = False) -> None:
    """Write a wrapped line block and move to the next line at the left margin."""
    safe_text = _break_long_words(_pdf_safe_text(text))
    if not safe_text.strip():
        pdf.ln(line_height / 2)
        return

    style = "B" if bold else ""
    pdf.set_font("Helvetica", style=style, size=font_size)
    pdf.multi_cell(
        pdf.epw,
        line_height,
        safe_text,
        new_x="LMARGIN",
        new_y="NEXT",
    )


def export_markdown(content: str, title: str = "DocuMind AI") -> tuple[bytes, str]:
    """
    Export content as a Markdown file.

    Returns:
        Tuple of (file bytes, suggested filename).
    """
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    header = f"# {title}\n\n_Gerado por DocuMind AI em {timestamp} UTC_\n\n"
    payload = (header + content.strip() + "\n").encode("utf-8")
    filename = f"{_sanitize_filename(title)}_{timestamp}.md"
    return payload, filename


def export_docx(content: str, title: str = "DocuMind AI") -> tuple[bytes, str]:
    """
    Export content as a Word document.

    Returns:
        Tuple of (file bytes, suggested filename).
    """
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    document.add_paragraph(f"Gerado por DocuMind AI em {timestamp}")

    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped:
            document.add_paragraph(stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    timestamp_file = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"{_sanitize_filename(title)}_{timestamp_file}.docx"
    return buffer.getvalue(), filename


def export_pdf(content: str, title: str = "DocuMind AI") -> tuple[bytes, str]:
    """
    Export content as a simple PDF document.

    Returns:
        Tuple of (file bytes, suggested filename).
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    _pdf_write_line(pdf, title, line_height=10, font_size=16, bold=True)
    pdf.ln(4)

    for line in content.splitlines():
        text = line.strip()
        if not text:
            pdf.ln(3)
            continue
        if text.startswith("## "):
            _pdf_write_line(pdf, text[3:].strip(), line_height=8, font_size=13, bold=True)
            continue
        if text.startswith("# "):
            _pdf_write_line(pdf, text[2:].strip(), line_height=8, font_size=14, bold=True)
            continue
        _pdf_write_line(pdf, text, line_height=6, font_size=11)

    timestamp_file = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"{_sanitize_filename(title)}_{timestamp_file}.pdf"
    output = pdf.output()
    if isinstance(output, str):
        output = output.encode("latin-1")
    return bytes(output), filename


def build_consolidated_report(
    *,
    collection_label: str,
    document_summaries: dict[str, str] | None = None,
    document_names: dict[str, str] | None = None,
    collection_summary: str | None = None,
    insights_dashboard: str | None = None,
    comparison_result: str | None = None,
    comparison_title: str | None = None,
    chat_messages: list[dict] | None = None,
) -> str:
    """Assemble a full session report for consolidated export."""
    sections: list[str] = [
        f"# Relatório DocuMind AI\n\n**Coleção:** {collection_label}\n",
    ]

    if collection_summary:
        sections.append("---\n\n## Resumo da coleção\n\n" + collection_summary.strip())

    if insights_dashboard:
        sections.append("---\n\n## Painel de Insights\n\n" + insights_dashboard.strip())

    if comparison_result:
        title = comparison_title or "Comparação de documentos"
        sections.append(f"---\n\n## {title}\n\n" + comparison_result.strip())

    if document_summaries:
        sections.append("---\n\n## Resumos por documento\n")
        for document_id, summary in document_summaries.items():
            label = (document_names or {}).get(document_id, document_id)
            sections.append(f"### {label}\n\n{summary.strip()}\n")

    if chat_messages:
        sections.append("---\n\n## Histórico de conversa\n")
        for message in chat_messages:
            role = message.get("role", "user")
            content = str(message.get("content", "")).strip()
            if not content:
                continue
            label = "Utilizador" if role == "user" else "Assistente"
            sections.append(f"**{label}:** {content}\n")
            if role == "assistant" and message.get("sources"):
                sections.append("\n_Fontes:_\n")
                for source in message["sources"]:
                    doc = source.get("document", "Documento")
                    page = source.get("page", "?")
                    excerpt = source.get("excerpt", "")
                    sections.append(f'- {doc} (pág. {page}): "{excerpt}"\n')
            sections.append("\n")

    return "\n".join(sections).strip() + "\n"

